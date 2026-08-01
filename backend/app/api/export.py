from io import BytesIO

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import Response
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import CurrentUser
from app.db import get_db
from app.models.models import Session, SessionDocument
from app.models.resume_schema import CoverLetterContent, ResumeContent
from app.services.latex.compiler import CompileError, LatexCompiler
from app.services.rendering.renderer import ResumeRenderer, tex_escape

router = APIRouter(prefix="/api/sessions", tags=["export"])


def _docx_apply_spans(paragraph, spans: list, text: str):
    """Apply span formatting (bold/italic/underline) as docx run properties."""
    if not spans:
        paragraph.add_run(text)
        return

    def _attr(span, key):
        if isinstance(span, dict):
            return span.get(key)
        return getattr(span, key, None)

    cursor = 0
    for span in sorted(spans, key=lambda s: _attr(s, "start")):
        s_start = _attr(span, "start")
        s_end = _attr(span, "end")
        if cursor < s_start:
            paragraph.add_run(text[cursor:s_start])
        run = paragraph.add_run(text[s_start:s_end])
        formats = _attr(span, "formats") or []
        if "bold" in formats:
            run.bold = True
        if "italic" in formats:
            run.italic = True
        if "underline" in formats:
            run.underline = True
        cursor = s_end
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _render_cover_letter_docx(content: CoverLetterContent) -> bytes:
    from docx import Document as DocxDocument

    docx = DocxDocument()
    if content.salutation:
        docx.add_paragraph(content.salutation)
    for para in content.paragraphs:
        docx.add_paragraph(para.text)
    if content.closing:
        docx.add_paragraph(content.closing)

    buffer = BytesIO()
    docx.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _render_cover_letter_pdf(content: CoverLetterContent) -> bytes:
    """Render cover letter as PDF via LaTeX."""
    lines = []
    if content.salutation:
        lines.append(tex_escape(content.salutation))
        lines.append("")
    for para in content.paragraphs:
        lines.append(tex_escape(para.text))
        lines.append("")
    if content.closing:
        lines.append(tex_escape(content.closing))

    tex = (
        r"""\documentclass[11pt,a4paper]{article}
\usepackage[margin=1in]{geometry}
\usepackage{parskip}
\begin{document}
"""
        + "\n".join(lines)
        + r"""
\end{document}"""
    )

    compiler = LatexCompiler()
    return compiler.compile(tex, "cover-letter")


@router.get("/{session_id}/export")
async def export_document(
    session_id: str,
    current_user: CurrentUser,
    db: AsyncSession = Depends(get_db),
    format: str = Query(pattern="^(tex|pdf|docx|txt|html)$"),
    doc_type: str = Query(default="resume"),
):
    result = await db.execute(
        select(Session).where(Session.id == session_id, Session.user_id == current_user.id)
    )
    session = result.scalar_one_or_none()
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    # Cover letter export
    if doc_type == "cover_letter":
        if format not in ("pdf", "docx"):
            raise HTTPException(
                status_code=400, detail="Cover letter export supports pdf and docx only"
            )

        doc_result = await db.execute(
            select(SessionDocument)
            .where(
                SessionDocument.session_id == session.id, SessionDocument.doc_type == "cover_letter"
            )
            .order_by(SessionDocument.version.desc())
            .limit(1)
        )
        doc = doc_result.scalar_one_or_none()
        if not doc:
            raise HTTPException(status_code=404, detail="No cover letter found")

        content_dict = doc.content_json or {}
        if "paragraphs" not in content_dict:
            cl_content = CoverLetterContent.from_legacy_text(content_dict.get("text", ""))
        else:
            cl_content = CoverLetterContent.model_validate(content_dict)

        if format == "docx":
            docx_bytes = _render_cover_letter_docx(cl_content)
            return Response(
                content=docx_bytes,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=cover_letter.docx"},
            )

        if format == "pdf":
            try:
                pdf_bytes = _render_cover_letter_pdf(cl_content)
                return Response(
                    content=pdf_bytes,
                    media_type="application/pdf",
                    headers={"Content-Disposition": "attachment; filename=cover_letter.pdf"},
                )
            except CompileError as e:
                raise HTTPException(status_code=400, detail=e.message)

    # Resume export (existing logic)
    doc_result = await db.execute(
        select(SessionDocument)
        .where(SessionDocument.session_id == session.id, SessionDocument.doc_type == "resume")
        .order_by(SessionDocument.version.desc())
        .limit(1)
    )
    doc = doc_result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="No document found")

    content = ResumeContent.model_validate(doc.content_json)
    renderer = ResumeRenderer()

    if format == "tex":
        tex = renderer.render_tex(content)
        return Response(
            content=tex,
            media_type="text/plain",
            headers={"Content-Disposition": "attachment; filename=resume.tex"},
        )

    if format == "html":
        html = renderer.render_html(content)
        return Response(
            content=html,
            media_type="text/html",
            headers={"Content-Disposition": "attachment; filename=resume.html"},
        )

    if format == "pdf":
        try:
            tex_to_compile = renderer.render_tex(content)
            compiler = LatexCompiler()
            pdf_bytes = compiler.compile(tex_to_compile, str(doc.id))
            return Response(
                content=pdf_bytes,
                media_type="application/pdf",
                headers={"Content-Disposition": "attachment; filename=resume.pdf"},
            )
        except CompileError as e:
            raise HTTPException(status_code=400, detail=e.message)

    if format == "docx":
        from docx import Document as DocxDocument

        docx = DocxDocument()
        if content.basics.name:
            docx.add_heading(content.basics.name, 0)
            info_parts = []
            if content.basics.email:
                info_parts.append(content.basics.email)
            if content.basics.phone:
                info_parts.append(content.basics.phone)
            if content.basics.location:
                info_parts.append(content.basics.location)
            if info_parts:
                p = docx.add_paragraph()
                p.add_run("  |  ".join(info_parts)).italic = True
        for section in content.sections:
            docx.add_heading(section.label, level=1)
            for entry in section.entries:
                p = docx.add_paragraph()
                run = p.add_run(entry.title)
                run.bold = True
                if entry.dates:
                    p.add_run(f"  —  {entry.dates}")
                if entry.role:
                    role_p = docx.add_paragraph()
                    role_p.add_run(entry.role).italic = True
                    if entry.location:
                        role_p.add_run(f"  —  {entry.location}").italic = True
                for bullet in entry.bullets:
                    bp = docx.add_paragraph(style="List Bullet")
                    _docx_apply_spans(
                        bp,
                        [s.model_dump() if hasattr(s, "model_dump") else s for s in bullet.spans],
                        bullet.text,
                    )
            for sk in section.skill_rows:
                p = docx.add_paragraph()
                run = p.add_run(f"{sk.category} ")
                run.bold = True
                p.add_run(sk.items)

        buffer = BytesIO()
        docx.save(buffer)
        buffer.seek(0)
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=resume.docx"},
        )

    if format == "txt":
        lines: list[str] = []
        if content.basics.name:
            lines.append(content.basics.name.upper())
            info_parts = []
            if content.basics.email:
                info_parts.append(content.basics.email)
            if content.basics.phone:
                info_parts.append(content.basics.phone)
            if content.basics.location:
                info_parts.append(content.basics.location)
            if info_parts:
                lines.append("  ".join(info_parts))
            lines.append("")
        for section in content.sections:
            lines.append(section.label.upper())
            lines.append("")
            for entry in section.entries:
                title_line = entry.title
                if entry.dates:
                    title_line += f"  --  {entry.dates}"
                lines.append(title_line)
                if entry.role:
                    role_line = f"    {entry.role}"
                    if entry.location:
                        role_line += f"  --  {entry.location}"
                    lines.append(role_line)
                for bullet in entry.bullets:
                    lines.append(f"  \u2022 {bullet.text}")
            for sk in section.skill_rows:
                lines.append(f"{sk.category}: {sk.items}")
            lines.append("")

        return Response(
            content="\n".join(lines),
            media_type="text/plain",
            headers={"Content-Disposition": "attachment; filename=resume.txt"},
        )

    raise HTTPException(status_code=400, detail="Invalid format")
